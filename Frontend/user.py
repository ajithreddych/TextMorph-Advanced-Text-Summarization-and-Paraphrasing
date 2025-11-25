import streamlit as st
import requests
import PyPDF2
import docx
import io
import pandas as pd
import math
import numpy as np
import matplotlib.pyplot as plt
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from rouge_score import rouge_scorer
import textstat
from datetime import datetime, date
import warnings
warnings.filterwarnings("ignore", message="pkg_resources is deprecated")


API_BASE = "http://127.0.0.1:5000"  # Backend root (Flask app)
API_SUMMARIZATION = f"{API_BASE}/api/summarize"  # Summarization endpoint
API_PARAPHRASE = f"{API_BASE}/api/paraphrase"  # Paraphrase generation endpoint

def user_dashboard():
    import re  # local import so we don't need to modify the file-level imports

    menu_label = st.sidebar.radio(
        "Navigation",
        ["📝 Summarization", "🔄 Paraphrasing", "👤 Profile", "📜 History", "📂 Dataset Evaluation", "🚪 Logout"]
    )

    # normalize label: remove emojis / non-word characters so comparisons are stable
    normalized = re.sub(r"[^\w\s]", "", menu_label).strip()
    # get the first word (e.g., "Summarization", "Paraphrasing", "Profile", "History", "Dataset", "Logout")
    menu_key = normalized.split()[0] if normalized else menu_label

    if menu_key == "Summarization":
        render_summarization_section()

    elif menu_key == "Paraphrasing":
        render_paraphrasing_section()

    elif menu_key == "History":
        render_history_section()

    elif menu_key == "Profile":
        render_profile_section()
    
    elif menu_key == "Dataset":
        render_dataset_evaluation_section()

    


    elif menu_key == "Logout":
        st.markdown("### 👋 Confirm Logout")
        st.info("Are you sure you want to log out?")
        if st.button("Logout Now", use_container_width=True):
            st.session_state.logged_in = False
            st.session_state.page = "login"
            st.session_state.token = None
            st.success("✅ Logged out successfully!")
            st.rerun()

    else:
        st.write(f"Selected User Menu: {menu_label}")
        st.info("User functionalities go here.")


# safe rerun helper — use instead of calling st.experimental_rerun() directly
def safe_rerun():
    """
    Try to use st.experimental_rerun(); if it does not exist in this Streamlit build,
    flip a session_state toggle which also triggers a rerun.
    """
    try:
        if hasattr(st, "experimental_rerun"):
            # some builds provide this; try-catch in case it errors
            try:
                st.experimental_rerun()
                return
            except Exception:
                pass
        # fallback: flipping a session_state key triggers a rerun
        st.session_state["_rerun_toggle"] = not st.session_state.get("_rerun_toggle", False)
    except Exception:
        # Last-resort no-op (avoid bringing down the app)
        pass



# -------------------- SUMMARIZATION SECTION --------------------
def render_summarization_section():
    st.title("📝 AI Summarization")
    st.write("Upload a document or enter text below to generate and evaluate summaries using fine-tuned models.")

    uploaded_file = st.file_uploader("📂 Upload a file (PDF, DOCX, TXT, CSV)", type=["pdf", "docx", "txt", "csv"])
    input_text = ""

    # -------- File Handling --------
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()
        try:
            if file_type == "pdf":
                reader = PyPDF2.PdfReader(uploaded_file)
                input_text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
            elif file_type == "docx":
                doc = docx.Document(uploaded_file)
                input_text = "\n".join([para.text for para in doc.paragraphs])
            elif file_type == "txt":
                input_text = uploaded_file.read().decode("utf-8")
            elif file_type == "csv":
                df = pd.read_csv(uploaded_file)
                if "text" in df.columns:
                    input_text = " ".join(df["text"].astype(str).head(20).tolist())
                else:
                    st.warning("No 'text' column found in CSV. Showing first 20 cells from first column.")
                    input_text = " ".join(df.iloc[:, 0].astype(str).head(20).tolist())
        except Exception as e:
            st.error(f"Error reading file: {e}")
    else:
        input_text = st.text_area("✏️ Enter or paste text here:", height=200)

    # -------- Reference Text Input --------
    st.subheader("📘 Optional Reference Text")
    reference_text = st.text_area("Enter reference summary (for evaluation):", height=150, placeholder="Paste the human-written or reference summary here...")

    # -------- Model Selection --------
    model_choice = st.selectbox("🤖 Select Model", ["t5-small", "flan-t5-small", "bart-base"])
    length_choice = st.radio("📏 Summary Length", ["short", "medium", "long"], horizontal=True)
    # ✅ Ensure previous summary persists across reruns
    if "generated_summary" not in st.session_state:
        st.session_state.generated_summary = ""
    if "translated_text" not in st.session_state:
        st.session_state.translated_text = ""
    # Initialize feedback-related session state to persist choices across reruns
    if "feedback_choice" not in st.session_state:
        st.session_state.feedback_choice = None
    if "feedback_comment" not in st.session_state:
        st.session_state.feedback_comment = ""
    generate_clicked = st.button("🚀 Generate Summary")

    # -------- Generate Summary Button --------
    if generate_clicked or st.session_state.generated_summary.strip():
        if not input_text.strip():
            st.warning("Please upload or enter text to summarize.")
        else:
            with st.spinner("Generating summary... please wait ⏳"):
                try:
                    payload = {"text": input_text, "model_name": model_choice, "length": length_choice}
                    res = requests.post(API_SUMMARIZATION, json=payload)
                    if res.status_code == 200:
                        data = res.json()
                        summary = data.get("summary", "").strip()
                        st.session_state.generated_summary = summary
                        st.success("✅ Summary generated successfully!")

                        # -------- Display in 3 Side-by-Side Boxes --------
                        col1, col2, col3 = st.columns(3)
                        with col1:
                            st.markdown("### 🧾 Input Text")
                            st.text_area("Original Text", input_text, height=300)
                        with col2:
                            st.markdown("### 🤖 Generated Summary")
                            st.text_area("Model Output", summary, height=300)
                        with col3:
                            st.markdown("### 📘 Reference Summary")
                            st.text_area("Reference Text", reference_text, height=300)
                        

                        # -------- Feedback Section --------
                        st.markdown("### 💬 Provide Feedback on this Summary")

                        # Thumbs buttons (store selection in session_state instead of auto-posting)
                        col_fb1, col_fb2 = st.columns([1, 4])
                        with col_fb1:
                            # Use buttons to set a persistent selection
                            if st.button("👍", key="thumbs_up"):
                                st.session_state.feedback_choice = "up"
                            if st.button("👎", key="thumbs_down"):
                                st.session_state.feedback_choice = "down"

                        with col_fb2:
                            # Show current selection to the user
                            current_sel = st.session_state.feedback_choice
                            if current_sel:
                                st.info(f"Selected rating: {'👍' if current_sel == 'up' else '👎'}")
                            else:
                                st.info("No rating selected yet. Click 👍 or 👎 to choose a rating.")

                        # Comment input (store in session_state to persist across reruns)
                        st.text_area("🗣️ Optional Comment", placeholder="What did you like or dislike?", key="feedback_comment", height=120)

                        # Submit Feedback button placed directly below the comment box
                        if st.button("💬 Submit Feedback", use_container_width=True, key="submit_feedback"):
                            # Read persistent values
                            feedback_choice = st.session_state.get("feedback_choice")
                            comment = st.session_state.get("feedback_comment", "").strip()

                            if not feedback_choice:
                                st.error("Please select 👍 or 👎 before submitting feedback.")
                            else:
                                try:
                                    headers = {"Authorization": st.session_state.token}
                                    feedback_payload = {
                                        "user_id": st.session_state.user_id,
                                        "task_type": "summarization",
                                        "task_id": st.session_state.get("last_summary_id", 0),
                                        "rating": feedback_choice,
                                        "comment": comment
                                    }
                                    res = requests.post(f"{API_BASE}/feedback/add", json=feedback_payload, headers=headers)
                                    if res.status_code == 201:
                                        st.success("✅ Thank you for your feedback!")
                                        # Safely attempt to clear only non-widget state. If clearing raises, ignore it.
                                        try:
                                            # feedback_choice is not tied to a widget, safe to reset; guard anyway
                                            st.session_state.feedback_choice = None
                                        except Exception:
                                            pass
                                        # DO NOT force-reset feedback_comment (widget). If you want to attempt, do it inside try/except:
                                        try:
                                            st.session_state.feedback_comment = ""
                                        except Exception:
                                            # ignore Streamlit error about modifying widget-backed state after instantiation
                                            pass
                                    else:
                                        # Mirror backend error message if provided
                                        try:
                                            err = res.json().get("error") or res.json().get("message") or "Failed to save feedback."
                                        except Exception:
                                            err = "Failed to save feedback."
                                        st.error(f"Failed to save feedback. {err}")
                                except Exception as e:
                                    st.error(f"Error submitting feedback: {e}")


                        

                        # --- Auto-save summary to DB ---
                        try:
                            headers = {"Authorization": st.session_state.token}
                            save_payload = {
                                "user_id": st.session_state.user_id,
                                "input_text": input_text,
                                "summary_text": summary,
                                "model_name": model_choice,
                                "size": length_choice.capitalize()
                            }
                            save_res = requests.post(f"{API_BASE}/feedback/summary/save", json=save_payload, headers=headers)
                            if save_res.status_code == 201:
                                summary_id = save_res.json().get("summary_id")
                                st.session_state.last_summary_id = summary_id
                                st.info("📥 Summary auto-saved successfully!")
                            else:
                                st.warning("⚠️ Could not auto-save summary.")
                        except Exception as e:
                            st.error(f"Auto-save error: {e}")


                        # -------- Evaluation Metrics --------
                        if reference_text.strip():
                            st.subheader("📊 Evaluation Metrics")
                            results = evaluate_summary(summary, reference_text)

                            bleu = results["BLEU"]
                            perplexity = results["Perplexity (approx)"]
                            readability_delta = abs(results["Readability Δ (Ref - Gen)"])
                            rouge1 = results["ROUGE-1 (F)"]
                            rouge2 = results["ROUGE-2 (F)"]
                            rougel = results["ROUGE-L (F)"]

                            # Horizontal metrics layout
                            metric_cols = st.columns(6)
                            metric_cols[0].metric("BLEU", f"{bleu:.3f}")
                            metric_cols[1].metric("Perplexity", f"{perplexity:.3f}")
                            metric_cols[2].metric("Readability Δ", f"{readability_delta:.3f}")
                            metric_cols[3].metric("ROUGE-1 (F)", f"{rouge1:.3f}")
                            metric_cols[4].metric("ROUGE-2 (F)", f"{rouge2:.3f}")
                            metric_cols[5].metric("ROUGE-L (F)", f"{rougel:.3f}")

                            # -------- Radar Chart Visualization --------
                            categories = ['BLEU', 'Perplexity', 'Readability Δ', 'ROUGE-1', 'ROUGE-2', 'ROUGE-L']
                            values = [bleu, perplexity, readability_delta, rouge1, rouge2, rougel]
                            values_ref = [0, 0, 0, 0, 0, 0]

                            norm_values = np.array(values)
                            norm_values = norm_values / (norm_values.max() if norm_values.max() != 0 else 1)

                            N = len(categories)
                            angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
                            norm_values = np.concatenate((norm_values, [norm_values[0]]))
                            angles += angles[:1]

                            fig, ax = plt.subplots(subplot_kw=dict(polar=True))
                            ax.plot(angles, norm_values, color='b', linewidth=2, label='Generated')
                            ax.fill(angles, norm_values, color='b', alpha=0.25)
                            ax.plot(angles, np.concatenate((values_ref, [values_ref[0]])), color='r', linewidth=2, label='Reference')
                            ax.set_yticklabels([])
                            ax.set_xticks(angles[:-1])
                            ax.set_xticklabels(categories, fontsize=10)
                            ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))
                            st.pyplot(fig)

                        else:
                            st.info("ℹ️ Add a reference text above to evaluate summary quality.")

                        # -------- Save / Download --------
                        col4, col5 = st.columns(2)
                        with col4:
                            st.download_button(
                                label="💾 Download Summary",
                                data=summary,
                                file_name="summary.txt",
                                mime="text/plain"
                            )
                        # Note: Removed the old "Save to History" button per request.
                    else:
                        st.error(res.json().get("error", "Failed to summarize text."))
                except Exception as e:
                    st.error(f"Error contacting server: {e}")

        # -------- TRANSLATION FEATURE (fixed persistent version) --------
        if st.session_state.generated_summary.strip():
            st.markdown("---")
            st.subheader("🌍 Translate Generated Summary")

            # ✅ Store the generated summary in session state to persist across reruns
            summary = st.session_state.generated_summary


            lang_map = {
                "Hindi": "hi",
                "Telugu": "te",
                "Tamil": "ta",
                "Kannada": "kn",
                "Malayalam": "ml",
                "French": "fr",
                "Spanish": "es",
                "German": "de",
                "Japanese": "ja"
            }

            target_lang = st.selectbox("Select Target Language:", list(lang_map.keys()), key="translate_lang")

            if st.button("🌐 Translate Summary", key="translate_btn"):
                if "generated_summary" not in st.session_state or not st.session_state.generated_summary.strip():
                    st.warning("Please generate a summary first before translating.")
                else:
                    with st.spinner("Translating summary... please wait ⏳"):
                        try:
                            res = requests.post(
                                f"{API_BASE}/api/translate",
                                json={
                                    "text": st.session_state.generated_summary,
                                    "target_lang": lang_map[target_lang]
                                }
                            )
                            if res.status_code == 200:
                                translated_text = res.json().get("translated_text", "")
                                st.session_state.translated_text = translated_text  # ✅ keep translation persistent
                                st.success(f"✅ Translated to {target_lang}")

                            else:
                                st.error("Translation failed. Please check backend or try again.")
                        except Exception as e:
                            st.error(f"Error during translation: {e}")

            # ✅ Show the translation if it exists
            if "translated_text" in st.session_state and st.session_state.translated_text.strip():
                st.text_area(f"🗣️ {target_lang} Translation", st.session_state.translated_text, height=200)
                st.download_button(
                    label=f"💾 Download {target_lang} Translation",
                    data=st.session_state.translated_text,
                    file_name=f"summary_{lang_map[target_lang]}.txt",
                    mime="text/plain"
                )

def render_paraphrasing_section():
    st.title("🔄 AI Paraphrasing")
    st.write("Upload a document or enter text below to generate paraphrases using fine-tuned models.")

    # API for paraphrase (backend route you created)
    API_PARAPHRASE = f"{API_BASE}/api/paraphrase"
    API_SAVE_PARAPHRASE = f"{API_BASE}/feedback/paraphrase/save"  # backend save route
    API_ADD_FEEDBACK = f"{API_BASE}/feedback/add"
    API_TRANSLATE = f"{API_BASE}/api/translate"  # translation endpoint

    uploaded_file = st.file_uploader(
        "📂 Upload a file (PDF, DOCX, TXT, CSV)",
        type=["pdf", "docx", "txt", "csv"],
        key="par_uploader"
    )
    input_text = ""

    # -------- File Handling (same logic as summarization) --------
    if uploaded_file:
        file_type = uploaded_file.name.split(".")[-1].lower()
        try:
            if file_type == "pdf":
                reader = PyPDF2.PdfReader(uploaded_file)
                input_text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
            elif file_type == "docx":
                doc = docx.Document(uploaded_file)
                input_text = "\n".join([para.text for para in doc.paragraphs])
            elif file_type == "txt":
                input_text = uploaded_file.read().decode("utf-8")
            elif file_type == "csv":
                df = pd.read_csv(uploaded_file)
                if "text" in df.columns:
                    input_text = " ".join(df["text"].astype(str).head(20).tolist())
                else:
                    st.warning("No 'text' column found in CSV. Showing first 20 cells from first column.")
                    input_text = " ".join(df.iloc[:, 0].astype(str).head(20).tolist())
        except Exception as e:
            st.error(f"Error reading file: {e}")
    else:
        input_text = st.text_area("✏️ Enter or paste text here:", height=200, key="paraphrase_input_text")

    # -------- Reference Text Input --------
    st.subheader("📘 Optional Reference Text")
    reference_text = st.text_area(
        "Enter reference paraphrase (for evaluation/comparison):",
        height=150,
        placeholder="Paste the human-written or reference paraphrase here...",
        key="paraphrase_reference"
    )

    # -------- Model Selection --------
    st.subheader("🤖 Select Paraphrase Model")
    model_map = {
        "t5-small": "training_pipeline/datasets/processed/model_outputs/paraphrase_t5-small_20251105_062028",
        "flan-t5-small": "training_pipeline/datasets/processed/model_outputs/paraphrase_google_flan-t5-small_20251105_065026",
        "bart-base": "training_pipeline/datasets/processed/model_outputs/paraphrase_facebook_bart-base_20251105_071833"
    }
    model_choice = st.selectbox("Model", list(model_map.keys()), index=0, help="Choose the fine-tuned paraphrase model", key="par_model_choice")

    # -------- Complexity / Tone --------
    st.subheader("⚙️ Paraphrase Style")
    complexity_choice = st.radio("Select Style", ["Creative", "Standard", "Basic"], horizontal=True, index=1, key="par_complexity")

    # Persist generated paraphrase & translated paraphrase across reruns
    st.session_state.setdefault("generated_paraphrase", "")
    st.session_state.setdefault("paraphrase_saved_id", None)
    st.session_state.setdefault("par_translated_text", "")  # NEW: store translation

    # Feedback state for paraphrase (persist, initialize before widget creation)
    st.session_state.setdefault("par_feedback_choice", None)     # "up" or "down"
    st.session_state.setdefault("par_feedback_comment", "")     # widget-backed

    # Generate button (only trigger generation when clicked)
    generate_clicked = st.button("🔁 Generate Paraphrase", key="generate_paraphrase_btn")
    newly_generated = False

    # -------- Generation: only run when button is clicked --------
    if generate_clicked:
        if not input_text.strip():
            st.warning("Please upload or enter text to paraphrase.")
        else:
            with st.spinner("Generating paraphrase... please wait ⏳"):
                try:
                    payload = {
                        "text": input_text,
                        "model_name": model_choice,
                        "complexity": complexity_choice.lower()
                    }
                    resp = requests.post(API_PARAPHRASE, json=payload)
                    if resp.status_code == 200:
                        data = resp.json()
                        paraphrase = data.get("paraphrase", "").strip() or data.get("output", "").strip()
                        if not paraphrase:
                            st.error("No paraphrase returned by server.")
                        else:
                            st.session_state.generated_paraphrase = paraphrase
                            # reset any previous translation when new paraphrase generated
                            st.session_state.par_translated_text = ""
                            newly_generated = True
                            st.success("✅ Paraphrase generated successfully!")

                            # Auto-save paraphrase to DB (only immediately after generation) - if user_id exists
                            if st.session_state.get("user_id") and st.session_state.get("token"):
                                try:
                                    headers = {"Authorization": st.session_state.token}
                                    save_payload = {
                                        "user_id": st.session_state.user_id,
                                        "input_text": input_text,
                                        "paraphrase_text": paraphrase,
                                        "model_name": model_choice,
                                        "size": complexity_choice,            # stored into size ENUM
                                        "complexity": complexity_choice       # stored into complexity ENUM
                                    }

                                    save_res = requests.post(API_SAVE_PARAPHRASE, json=save_payload, headers=headers)
                                    if save_res.status_code == 201:
                                        paraphrase_id = save_res.json().get("paraphrase_id")
                                        st.session_state.paraphrase_saved_id = paraphrase_id
                                        st.info("📥 Paraphrase auto-saved successfully!")
                                    else:
                                        st.warning("⚠️ Could not auto-save paraphrase. You can save manually below.")
                                except Exception as e:
                                    st.warning(f"Auto-save error: {e}")
                            else:
                                st.info("Login required to auto-save paraphrase. Use manual Save button below.")
                    else:
                        try:
                            err = resp.json().get("error") or resp.json().get("message")
                        except Exception:
                            err = "Failed to paraphrase text."
                        st.error(err)
                except Exception as e:
                    st.error(f"Error contacting paraphrase server: {e}")

    # -------- Display generated paraphrase and evaluation (do NOT re-trigger generation) --------
    if st.session_state.generated_paraphrase.strip():
        paraphrase = st.session_state.generated_paraphrase

        # -------- Display in 3 Side-by-Side Boxes --------
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown("### 🧾 Input Text")
            st.text_area("Original Text", input_text, height=300, key="par_input_display")
        with c2:
            st.markdown("### ✍️ Generated Paraphrase")
            st.text_area("Paraphrase Output", paraphrase, height=300, key="par_output_display")
        with c3:
            st.markdown("### 📘 Reference Paraphrase")
            st.text_area("Reference Text", reference_text, height=300, key="par_reference_display")

        # -------- Provide Download & Manual Save --------
        col_dl, col_save = st.columns(2)
        with col_dl:
            st.download_button(
                label="💾 Download Paraphrase",
                data=paraphrase,
                file_name="paraphrase.txt",
                mime="text/plain",
                key="par_download"
            )

        with col_save:
            if st.button("💾 Save Paraphrase (Manual)", key="manual_save_paraphrase"):
                if not st.session_state.get("user_id") or not st.session_state.get("token"):
                    st.error("You must be logged in to save paraphrase.")
                else:
                    try:
                        headers = {"Authorization": st.session_state.token}
                        save_payload = {
                            "user_id": st.session_state.user_id,
                            "input_text": input_text,
                            "paraphrase_text": paraphrase,
                            "model_name": model_choice,
                            "size": complexity_choice
                        }
                        save_res = requests.post(API_SAVE_PARAPHRASE, json=save_payload, headers=headers)
                        if save_res.status_code == 201:
                            paraphrase_id = save_res.json().get("paraphrase_id")
                            st.session_state.paraphrase_saved_id = paraphrase_id
                            st.success("✅ Paraphrase saved to your history!")
                        else:
                            try:
                                err = save_res.json().get("error") or save_res.json().get("message")
                            except Exception:
                                err = "Save failed."
                            st.error(f"Failed to save paraphrase. {err}")
                    except Exception as e:
                        st.error(f"Error saving paraphrase: {e}")

        # -------- Translate Paraphrase UI (NEW) --------
        st.markdown("---")
        st.subheader("🌍 Translate Generated Paraphrase")

        # language map same as summarization
        lang_map = {
            "Hindi": "hi",
            "Telugu": "te",
            "Tamil": "ta",
            "Kannada": "kn",
            "Malayalam": "ml",
            "French": "fr",
            "Spanish": "es",
            "German": "de",
            "Japanese": "ja",
            "English": "en"
        }

        # Persist selected target language (avoid collision)
        target_lang = st.selectbox("Select Target Language:", list(lang_map.keys()), key="par_translate_lang")

        # Translate button
        if st.button("🌐 Translate Paraphrase", key="par_translate_btn"):
            if not st.session_state.get("generated_paraphrase"):
                st.warning("Please generate a paraphrase first before translating.")
            else:
                with st.spinner("Translating paraphrase... please wait ⏳"):
                    try:
                        res = requests.post(
                            API_TRANSLATE,
                            json={
                                "text": st.session_state.generated_paraphrase,
                                "target_lang": lang_map[target_lang]
                            }
                        )
                        if res.status_code == 200:
                            translated_text = res.json().get("translated_text", "")
                            st.session_state.par_translated_text = translated_text or ""
                            st.success(f"✅ Translated to {target_lang}")
                        else:
                            try:
                                err = res.json().get("error") or res.json().get("message") or "Translation failed."
                            except Exception:
                                err = "Translation failed."
                            st.error(err)
                    except Exception as e:
                        st.error(f"Error during translation: {e}")

        # show translated text + download if exists
        if st.session_state.get("par_translated_text", "").strip():
            st.text_area(f"🗣️ {target_lang} Translation", st.session_state.par_translated_text, height=200, key="par_translated_display")
            st.download_button(
                label=f"💾 Download {target_lang} Translation",
                data=st.session_state.par_translated_text,
                file_name=f"paraphrase_{lang_map[target_lang]}.txt",
                mime="text/plain",
                key="par_translated_download"
            )

        st.markdown("---")

        # -------- Feedback Section for Paraphrase (thumbs + comment) --------
        st.markdown("### 💬 Provide Feedback on this Paraphrase")
        fb_col1, fb_col2 = st.columns([1, 4])
        with fb_col1:
            if st.button("👍", key="par_thumbs_up"):
                st.session_state.par_feedback_choice = "up"
            if st.button("👎", key="par_thumbs_down"):
                st.session_state.par_feedback_choice = "down"
        with fb_col2:
            current_par_sel = st.session_state.get("par_feedback_choice")
            if current_par_sel:
                st.info(f"Selected rating: {'👍' if current_par_sel == 'up' else '👎'}")
            else:
                st.info("No rating selected yet. Click 👍 or 👎 to choose a rating.")

        # comment widget (must use a key and avoid modifying this key after creation)
        st.text_area("🗣️ Optional Comment", placeholder="What did you like or dislike?", key="par_feedback_comment", height=100)

        # submit feedback button
        if st.button("💬 Submit Feedback", use_container_width=True, key="par_submit_feedback"):
            par_feedback_choice = st.session_state.get("par_feedback_choice")
            par_comment = st.session_state.get("par_feedback_comment", "").strip()
            if not par_feedback_choice:
                st.error("Please select 👍 or 👎 before submitting feedback.")
            else:
                try:
                    headers = {"Authorization": st.session_state.token} if st.session_state.get("token") else {}
                    feedback_payload = {
                        "user_id": st.session_state.get("user_id"),
                        "task_type": "paraphrase",
                        "task_id": st.session_state.get("paraphrase_saved_id", 0) or 0,
                        "rating": par_feedback_choice,
                        "comment": par_comment
                    }
                    res = requests.post(API_ADD_FEEDBACK, json=feedback_payload, headers=headers)
                    if res.status_code == 201:
                        st.success("✅ Thank you for your feedback!")
                        # clear only non-widget state safely
                        try:
                            st.session_state.par_feedback_choice = None
                        except Exception:
                            pass
                        # do not directly reassign par_feedback_comment (widget-backed) without try/except
                        try:
                            st.session_state.par_feedback_comment = ""
                        except Exception:
                            pass
                    else:
                        try:
                            err = res.json().get("error") or res.json().get("message") or "Failed to save feedback."
                        except Exception:
                            err = "Failed to save feedback."
                        st.error(f"Failed to save feedback. {err}")
                except Exception as e:
                    st.error(f"Error submitting feedback: {e}")

        # -------- Optional Evaluation (if reference provided) --------
        if reference_text.strip():
            st.subheader("📊 Paraphrase Comparison Metrics")
            try:
                results = evaluate_summary(paraphrase, reference_text)
                bleu = results["BLEU"]
                perplexity = results["Perplexity (approx)"]
                readability_delta = abs(results["Readability Δ (Ref - Gen)"])
                rouge1 = results["ROUGE-1 (F)"]
                rouge2 = results["ROUGE-2 (F)"]
                rougel = results["ROUGE-L (F)"]

                # Horizontal metrics layout (same as summarization)
                metric_cols = st.columns(6)
                metric_cols[0].metric("BLEU", f"{bleu:.3f}")
                metric_cols[1].metric("Perplexity", f"{perplexity:.3f}")
                metric_cols[2].metric("Readability Δ", f"{readability_delta:.3f}")
                metric_cols[3].metric("ROUGE-1 (F)", f"{rouge1:.3f}")
                metric_cols[4].metric("ROUGE-2 (F)", f"{rouge2:.3f}")
                metric_cols[5].metric("ROUGE-L (F)", f"{rougel:.3f}")

                norm_values = np.array([bleu, perplexity, readability_delta, rouge1, rouge2, rougel], dtype=float)
                max_val = norm_values.max() if norm_values.max() != 0 else 1.0
                norm_values = norm_values / max_val

                N = len(norm_values)
                angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
                norm_values = np.concatenate((norm_values, [norm_values[0]]))
                angles += angles[:1]

                fig, ax = plt.subplots(subplot_kw=dict(polar=True))
                ax.plot(angles, norm_values, linewidth=2, label='Paraphrase')
                ax.fill(angles, norm_values, alpha=0.25)
                ax.set_yticklabels([])
                ax.set_xticks(angles[:-1])
                ax.set_xticklabels(['BLEU','Perplexity','Readability Δ','ROUGE-1','ROUGE-2','ROUGE-L'], fontsize=10)
                ax.set_title("Paraphrase Evaluation Radar", y=1.08)
                ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1))
                st.pyplot(fig)

            except Exception as e:
                st.warning(f"Evaluation error: {e}")

def render_history_section():
    """
    UI for user History (Summaries / Paraphrases).
    Expects backend endpoints (preferred):
      GET /history/summaries
      GET /history/paraphrases
    Fallback candidates are checked if preferred routes are not present.
    """
    st.title("📜 My History")
    st.write("View your past Summaries and Paraphrases.")

    if not st.session_state.get("user_id"):
        st.warning("Please log in to view your history.")
        return

    headers = {}
    if st.session_state.get("token"):
        headers["Authorization"] = st.session_state.token

    # Choose which history to show
    hist_choice = st.radio("Select history to view:", ["Summaries", "Paraphrases"], horizontal=True, key="history_choice")

    # paging controls
    per_page = st.selectbox("Items per page:", [10, 20, 50, 100], index=1, key="history_per_page")
    page = st.number_input("Page", min_value=1, value=1, step=1, key="history_page")

    # endpoint candidates (preferred first)
    endpoints = {
        "Summaries": [
            f"{API_BASE}/history/summaries",
            f"{API_BASE}/feedback/summaries",          # fallback
            f"{API_BASE}/feedback/summary/list"       # fallback
        ],
        "Paraphrases": [
            f"{API_BASE}/history/paraphrases",
            f"{API_BASE}/feedback/paraphrases",       # fallback
            f"{API_BASE}/feedback/paraphrase/list"    # fallback
        ]
    }

    chosen = hist_choice
    candidates = endpoints.get(chosen, [])

    data = None
    last_error = None

    # Try each candidate until one returns 200
    for url in candidates:
        try:
            # Optionally pass user_id as param if backend expects it
            params = {"user_id": st.session_state.get("user_id")}
            resp = requests.get(url, headers=headers, params=params, timeout=10)
            if resp.status_code == 200:
                # Accept list of dicts
                try:
                    data = resp.json()
                except Exception:
                    data = None
                # If backend wraps in { "data": [...] }
                if isinstance(data, dict) and "data" in data and isinstance(data["data"], list):
                    data = data["data"]
                # if it's a dict keyed by 'history' or similarly, try common keys
                if isinstance(data, dict) and not isinstance(data, list):
                    for k in ("history", "items", "rows", "results"):
                        if k in data and isinstance(data[k], list):
                            data = data[k]
                            break
                # ensure it's a list now
                if isinstance(data, list):
                    break
                else:
                    last_error = f"Unexpected response shape from {url}"
                    data = None
            else:
                # record backend message if present
                try:
                    j = resp.json()
                    last_error = j.get("error") or j.get("message") or f"HTTP {resp.status_code} from {url}"
                except Exception:
                    last_error = f"HTTP {resp.status_code} from {url}"
        except Exception as e:
            last_error = str(e)
            data = None

    # Nothing worked
    if not data:
        if last_error:
            st.error(f"Could not fetch {chosen.lower()} history. {last_error}")
        else:
            st.info(f"No {chosen.lower()} history found.")
        return

    # Normalize records into a DataFrame-friendly list of dicts
    rows = []
    for row in data:
        # Support different backend field names; attempt common ones
        if not isinstance(row, dict):
            continue
        if chosen == "Summaries":
            rows.append({
                "id": row.get("summary_id") or row.get("id") or row.get("summaryId") or "",
                "input_text": row.get("input_text") or row.get("text") or row.get("source") or "",
                "output_text": row.get("summary_text") or row.get("summary") or row.get("output") or "",
                "model": row.get("model_name") or row.get("model") or "",
                "size": row.get("size") or row.get("length") or "",
                "created_at": row.get("created_at") or row.get("timestamp") or row.get("created") or ""
            })
        else:  # Paraphrases
            rows.append({
                "id": row.get("paraphrase_id") or row.get("id") or row.get("paraphraseId") or "",
                "input_text": row.get("input_text") or row.get("text") or row.get("source") or "",
                "output_text": row.get("paraphrase_text") or row.get("paraphrase") or row.get("output") or "",
                "model": row.get("model_name") or row.get("model") or "",
                "size": row.get("size") or row.get("complexity") or "",
                "created_at": row.get("created_at") or row.get("timestamp") or row.get("created") or ""
            })

    # Reverse sort by created_at if present (latest first)
    try:
        rows.sort(key=lambda r: r.get("created_at") or "", reverse=True)
    except Exception:
        pass

    total = len(rows)
    start = (page - 1) * per_page
    end = start + per_page
    page_rows = rows[start:end]

    st.markdown(f"**Showing {chosen} ({start+1} - {min(end, total)}) of {total}**")

    if not page_rows:
        st.info("No records on this page. Try a smaller page number.")
        return

    # Build a compact preview table
    preview = []
    for r in page_rows:
        preview.append({
            "ID": r["id"],
            "Input (preview)": (r["input_text"][:120] + "…") if r["input_text"] and len(r["input_text"]) > 120 else r["input_text"],
            "Output (preview)": (r["output_text"][:140] + "…") if r["output_text"] and len(r["output_text"]) > 140 else r["output_text"],
            "Model": r["model"],
            "Size/Style": r["size"],
            "Created": r["created_at"]
        })

    df = pd.DataFrame(preview)
    st.dataframe(df, use_container_width=True)

    # Detailed expanders per item (for download / full view)
    st.markdown("---")
    st.subheader("Detailed entries")
    for r in page_rows:
        title = f"{chosen[:-1]} ID: {r['id']} — {r.get('created_at','')}"
        with st.expander(title, expanded=False):
            st.write("**Model:**", r.get("model", ""))
            st.write("**Size / Style:**", r.get("size", ""))
            st.write("**Created At:**", r.get("created_at", ""))
            st.markdown("**Input Text:**")
            st.text_area("input", r.get("input_text", ""), height=220, key=f"input_{r['id']}")
            st.markdown("**Output Text:**")
            st.text_area("output", r.get("output_text", ""), height=220, key=f"output_{r['id']}")
            # Downloads
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="📥 Download Input",
                    data=r.get("input_text", ""),
                    file_name=f"{chosen.lower()}_{r['id']}_input.txt",
                    mime="text/plain",
                    key=f"dl_in_{r['id']}"
                )
            with col2:
                st.download_button(
                    label="📥 Download Output",
                    data=r.get("output_text", ""),
                    file_name=f"{chosen.lower()}_{r['id']}_output.txt",
                    mime="text/plain",
                    key=f"dl_out_{r['id']}"
                )
            # Quick actions (optional): copy text to clipboard or re-run (you can implement re-run later)
            st.write("---")
            st.info("You can copy the text above or download it. To re-run this text through the model, copy it into the Summarization/Paraphrasing tab.")

    # Pagination controls summary
    st.markdown("---")
    st.write(f"Page {page} — showing {len(page_rows)} records. Total records: {total}.")


def render_dataset_evaluation_section():
    st.title("📂 Dataset Evaluation")
    st.write("Upload a CSV dataset with source text and reference text. Select task & model, then click **Evaluate** to run a batch evaluation across the file.")

    API_DATASET_EVAL = f"{API_BASE}/dataset/evaluate"

    uploaded_file = st.file_uploader("Upload CSV file", type=["csv"], key="dataset_eval_uploader")
    if not uploaded_file:
        st.info("Upload a CSV to begin. CSV must contain a column with source text and a column with reference text.")
        return

    # Read CSV preview (do not assume huge files; use head for preview)
    try:
        df = pd.read_csv(uploaded_file)
    except Exception as e:
        st.error(f"Error reading CSV: {e}")
        return

    st.markdown("**Preview (first 5 rows)**")
    st.dataframe(df.head(5))

    # let user choose columns for input and reference
    cols = df.columns.tolist()
    if not cols:
        st.error("CSV has no columns.")
        return

    st.subheader("Configuration")
    task = st.radio("Task", ["Summarization", "Paraphrasing"], index=0, horizontal=True, key="dataset_task")
    input_col = st.selectbox("Source text column", cols, index=0, key="dataset_input_col")
    ref_col = st.selectbox("Reference text column (human summary/paraphrase)", cols, index=min(1, len(cols)-1), key="dataset_ref_col")

    # model choices (reuse same models as UI for single-run)
    summarization_models = ["t5-small", "flan-t5-small", "bart-base"]
    paraphrase_models = ["t5-small", "flan-t5-small", "bart-base"]

    if task == "Summarization":
        model_choice = st.selectbox("Model (for summarization)", summarization_models, index=0, key="dataset_model_choice")
        complexity_choice = None
    else:
        model_choice = st.selectbox("Model (for paraphrasing)", paraphrase_models, index=0, key="dataset_model_choice_par")
        complexity_choice = st.selectbox("Paraphrase style", ["Creative", "Standard", "Basic"], index=1, key="dataset_par_style")

    sample_size = st.number_input("Max rows to evaluate (0 = all)", min_value=0, value=200, step=50, help="Limit rows for faster evaluation")

    # Evaluate button
    if st.button("▶️ Evaluate Dataset", key="dataset_evaluate_btn"):
        # prepare payload and file
        with st.spinner("Uploading dataset and starting evaluation... this may take a while depending on model & dataset size ⏳"):
            try:
                # reset any previous results
                st.session_state.pop("dataset_eval_result", None)

                # prepare file bytes: re-read uploaded_file to get bytes (Streamlit keeps it in-memory)
                uploaded_file.seek(0)
                files = {
                    "file": (uploaded_file.name, uploaded_file.getvalue(), "text/csv")
                }
                data = {
                    "task": task.lower(),             # 'summarization' or 'paraphrasing'
                    "model_name": model_choice,
                    "input_col": input_col,
                    "ref_col": ref_col,
                    "sample_size": int(sample_size) if sample_size and sample_size > 0 else 0
                }
                # include complexity if paraphrase
                if complexity_choice:
                    data["complexity"] = complexity_choice.lower()

                # send request
                res = requests.post(API_DATASET_EVAL, files=files, data=data, headers={"Authorization": st.session_state.get("token", "")})
                if res.status_code != 200:
                    try:
                        err = res.json().get("error") or res.json().get("message")
                    except Exception:
                        err = res.text
                    st.error(f"Evaluation failed: {err}")
                else:
                    result = res.json()
                    st.session_state["dataset_eval_result"] = result
                    st.success("✅ Evaluation finished (results in this page).")
                    safe_rerun()

            except Exception as e:
                st.error(f"Error calling dataset evaluate endpoint: {e}")

    # If a result exists in session_state, show it
    result = st.session_state.get("dataset_eval_result")
    if not result:
        return

    agg = result.get("aggregate", {}) or {}
    per_sample = result.get("per_sample", []) or []

    st.markdown("---")
    st.header("Aggregate scores")
    # arrange metrics in columns
    c1, c2, c3 = st.columns(3)
    c1.write(f"**BLEU:** {agg.get('BLEU', agg.get('bleu', 0)):.3f}")
    c1.write(f"**ROUGE-1:** {agg.get('ROUGE-1', agg.get('rouge1', 0)):.3f}")
    c2.write(f"**ROUGE-2:** {agg.get('ROUGE-2', agg.get('rouge2', 0)):.3f}")
    c2.write(f"**Perplexity:** {agg.get('Perplexity', agg.get('perplexity', 0)):.3f}")
    c3.write(f"**ROUGE-L:** {agg.get('ROUGE-L', agg.get('rougel', 0)):.3f}")
    c3.write(f"**Readability Δ:** {agg.get('Readability Delta', agg.get('readability_delta', 0)):.3f}")

    # Radar chart (similar style to other charts in app)
    try:
        categories = ['BLEU', 'Perplexity', 'Readability Δ', 'ROUGE-1', 'ROUGE-2', 'ROUGE-L']
        values = [
            float(agg.get('BLEU', agg.get('bleu', 0))),
            float(agg.get('Perplexity', agg.get('perplexity', 0))),
            abs(float(agg.get('Readability Delta', agg.get('readability_delta', 0)))),  # absolute delta for plotting
            float(agg.get('ROUGE-1', agg.get('rouge1', 0))),
            float(agg.get('ROUGE-2', agg.get('rouge2', 0))),
            float(agg.get('ROUGE-L', agg.get('rougel', 0)))
        ]

        # normalize safely
        norm_values = np.array(values, dtype=float)
        max_val = norm_values.max() if norm_values.max() != 0 else 1.0
        norm_values = norm_values / max_val

        N = len(categories)
        angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
        norm_values = np.concatenate((norm_values, [norm_values[0]]))
        angles += angles[:1]

        fig, ax = plt.subplots(subplot_kw=dict(polar=True), figsize=(7, 6))
        ax.plot(angles, norm_values, linewidth=2, label='Generated')
        ax.fill(angles, norm_values, alpha=0.25)
        ax.set_yticklabels([])
        ax.set_xticks(angles[:-1])
        ax.set_xticklabels(categories, fontsize=10)
        ax.set_title("Aggregate Evaluation Radar", y=1.05)
        ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.1))
        st.pyplot(fig)
    except Exception as e:
        st.warning(f"Could not draw radar chart: {e}")

    st.markdown("---")
    st.header("Per-sample results (first 20)")
    if per_sample:
        # take small preview
        preview_df = pd.DataFrame(per_sample).head(20)
        # show only selected columns if present
        cols_to_show = [c for c in ["id", "input", "reference", "generated", "bleu", "rouge1", "rouge2", "rougel"] if c in preview_df.columns]
        if cols_to_show:
            st.dataframe(preview_df[cols_to_show])
        else:
            st.dataframe(preview_df)
    else:
        st.info("No per-sample data returned from backend.")

    # allow user to download full per-sample CSV if available
    if per_sample:
        try:
            full_df = pd.DataFrame(per_sample)
            csv_bytes = full_df.to_csv(index=False).encode("utf-8")
            st.download_button("📥 Download full per-sample results (CSV)", data=csv_bytes, file_name="dataset_evaluation_results.csv", mime="text/csv")
        except Exception as e:
            st.warning(f"Could not prepare download: {e}")

    st.markdown("#### Notes")
    st.write("- Large datasets and model selection may take a long time; consider limiting rows with **Max rows to evaluate**.")
    st.write("- Backend should return `aggregate` (dict) and `per_sample` (list) for the best UX.")



# -------------------- EVALUATION FUNCTIONS --------------------
def evaluate_summary(generated, reference):
    """Compute ROUGE, BLEU, Perplexity & Readability delta."""
    results = {}

    # ROUGE (1,2,L)
    scorer = rouge_scorer.RougeScorer(["rouge1", "rouge2", "rougeL"], use_stemmer=True)
    rouge_scores = scorer.score(reference, generated)
    results["ROUGE-1 (F)"] = round(rouge_scores["rouge1"].fmeasure, 4)
    results["ROUGE-2 (F)"] = round(rouge_scores["rouge2"].fmeasure, 4)
    results["ROUGE-L (F)"] = round(rouge_scores["rougeL"].fmeasure, 4)

    # BLEU
    smoothie = SmoothingFunction().method4
    bleu = sentence_bleu([reference.split()], generated.split(), smoothing_function=smoothie)
    results["BLEU"] = round(bleu, 4)

    # Perplexity (approx via entropy of word distribution)
    words = generated.split()
    freq = {w: words.count(w) / len(words) for w in set(words)}
    entropy = -sum(p * math.log(p, 2) for p in freq.values())
    results["Perplexity (approx)"] = round(math.pow(2, entropy), 2)

    # Readability delta (Flesch Reading Ease)
    ref_read = textstat.flesch_reading_ease(reference)
    gen_read = textstat.flesch_reading_ease(generated)
    results["Readability Δ (Ref - Gen)"] = round(ref_read - gen_read, 2)

    return results


# -------------------- HELPER: SAFE DATE PARSER --------------------
def parse_date(dob_value):
    if not dob_value:
        return date.today()
    try:
        return date.fromisoformat(str(dob_value))
    except Exception:
        try:
            return datetime.strptime(str(dob_value), "%a, %d %b %Y %H:%M:%S %Z").date()
        except Exception:
            return date.today()


# -------------------- PROFILE SECTION --------------------
def render_profile_section():
    st.title("👤 My Profile")
    headers = {"Authorization": st.session_state.token}

    try:
        res = requests.get(f"{API_BASE}/profile/view", headers=headers)
        data = res.json()
        if res.status_code == 200:
            user = data["user"]
        else:
            st.error(data.get("message", "Failed to load profile"))
            return
    except Exception as e:
        st.error(f"Error fetching profile: {e}")
        return

    with st.expander("📄 View Profile", expanded=True):
        st.write(f"**Name:** {user['name']}")
        st.write(f"**Email:** {user['email']}")
        st.write(f"**Age:** {user['age']}")
        st.write(f"**Date of Birth:** {user['dob']}")
        st.write(f"**Gender:** {user['gender']}")
        st.write(f"**Account Created:** {user['created_at']}")

    with st.expander("✏️ Edit Profile"):
        name = st.text_input("Full Name", value=user["name"])
        age = st.number_input("Age", 1, 120, value=user["age"] if user["age"] else 18)
        dob = st.date_input("Date of Birth", parse_date(user["dob"]),
                            min_value=date(1900, 1, 1), max_value=date.today())
        gender = st.selectbox(
            "Gender",
            ["Male", "Female", "Other"],
            index=["Male", "Female", "Other"].index(user["gender"])
            if user["gender"] in ["Male", "Female", "Other"] else 0
        )

        if st.button("Update Profile"):
            payload = {"name": name, "age": age, "dob": str(dob), "gender": gender}
            try:
                res = requests.put(f"{API_BASE}/profile/update", json=payload, headers=headers)
                data = res.json()
                if res.status_code == 200:
                    st.success("✅ Profile updated successfully!")
                    st.rerun()
                else:
                    st.error(data.get("message", "Update failed"))
            except Exception as e:
                st.error(f"Server error: {e}")

    # ---------- CHANGE PASSWORD ----------
    with st.expander("🔐 Change Password"):
        old_pass = st.text_input("Old Password", type="password")
        new_pass = st.text_input("New Password", type="password")
        confirm_pass = st.text_input("Confirm New Password", type="password")

        if st.button("Change Password"):
            if not all([old_pass, new_pass, confirm_pass]):
                st.error("Please fill all fields.")
            elif new_pass != confirm_pass:
                st.error("New passwords do not match.")
            else:
                payload = {"old_password": old_pass, "new_password": new_pass}
                try:
                    res = requests.put(f"{API_BASE}/profile/change_password", json=payload, headers=headers)
                    data = res.json()
                    if res.status_code == 200:
                        st.success("✅ Password changed successfully!")
                    else:
                        st.error(data.get("message", "Password change failed"))
                except Exception as e:
                    st.error(f"Server error: {e}")
    # ---------- DELETE ACCOUNT ----------
    with st.expander("⚠️ Delete Account"):
        st.warning("This will permanently delete your account and all related data.")
        confirm = st.checkbox("I confirm that I want to delete my account permanently")

        if st.button("Delete My Account"):
            if confirm:
                try:
                    res = requests.delete(f"{API_BASE}/profile/delete", headers=headers)
                    data = res.json()
                    if res.status_code == 200:
                        st.success("✅ Account deleted successfully!")
                        st.session_state.logged_in = False
                        st.session_state.page = "login"
                        st.session_state.token = None
                        st.rerun()
                    else:
                        st.error(data.get("message", "Account deletion failed"))
                except Exception as e:
                    st.error(f"Server error: {e}")
            else:
                st.info("Please confirm deletion before proceeding.")
